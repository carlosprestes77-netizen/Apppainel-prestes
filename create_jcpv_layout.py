from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import copy

# ── Paleta JCPV ──────────────────────────────────────────────────────────────
DARK      = '1E1E1E'
GOLD      = 'B89B6A'
DARK_RGB  = RGBColor(0x1E, 0x1E, 0x1E)
GOLD_RGB  = RGBColor(0xB8, 0x9B, 0x6A)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_RGB  = RGBColor(0x88, 0x88, 0x88)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    hex_color = hex_color.lstrip('#')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=0, left=0, bottom=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def set_row_height(row, twips, exact=True):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(twips))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    trPr.append(trHeight)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def set_table_width_pct(table, pct=100):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(pct * 50))
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

def no_space(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)

def add_run(para, text, font='Times New Roman', size=12, bold=False,
            italic=False, color=None, underline=False):
    r = para.add_run(text)
    r.font.name  = font
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.italic = italic
    r.font.underline = underline
    if color:
        r.font.color.rgb = color
    return r

def para_border_bottom(para, color='B89B6A', sz='12'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_line_spacing(para, spacing=1.5):
    from docx.shared import Pt
    para.paragraph_format.line_spacing = spacing * Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()
sec = doc.sections[0]

# Tamanho A4
sec.page_height = Cm(29.7)
sec.page_width  = Cm(21.0)

# Margens ABNT para petições
sec.top_margin    = Cm(3.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.0)

# Distância do cabeçalho/rodapé à borda da folha
sec.header_distance = Cm(1.0)
sec.footer_distance = Cm(1.0)

# Estilos base
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after  = Pt(0)

# ─────────────────────────────────────────────────────────────────────────────
# CABEÇALHO (Word nativo)
# ─────────────────────────────────────────────────────────────────────────────

header = sec.header
# Limpa parágrafo padrão
for p in header.paragraphs:
    p.clear()

# Tabela no cabeçalho: fundo escuro cobrindo toda a largura
ht = header.add_table(rows=1, cols=1, width=Cm(21))
no_borders(ht)
set_table_width_pct(ht, 100)

hc = ht.rows[0].cells[0]
set_cell_bg(hc, DARK)
set_cell_margins(hc, top=300, left=500, bottom=220, right=500)
set_row_height(ht.rows[0], 1500, exact=False)

# Remove o parágrafo vazio que sobrou acima da tabela
for p in header.paragraphs:
    p._element.getparent().remove(p._element)

# Logo JCPV como imagem
p_logo = hc.paragraphs[0]
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p_logo)
p_logo.paragraph_format.space_before = Pt(6)
p_logo.paragraph_format.space_after  = Pt(4)
run_img = p_logo.add_run()
run_img.add_picture('/home/user/Apppainel-prestes/jcpv_logo.png', height=Cm(3.2))

# Slogan
p_slg = hc.add_paragraph()
p_slg.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p_slg)
p_slg.paragraph_format.space_before = Pt(2)
r = add_run(p_slg, 'A defesa técnica do seu patrimônio.',
            font='Garamond', size=8, italic=True,
            color=RGBColor(0xCC, 0xBB, 0x99))

# Linha dourada separadora dentro da célula
p_sep = hc.add_paragraph()
no_space(p_sep)
p_sep.paragraph_format.space_before = Pt(6)
para_border_bottom(p_sep, color=GOLD, sz='8')

# Info: OAB | Sigilo | Nacional
p_cred = hc.add_paragraph()
p_cred.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p_cred)
p_cred.paragraph_format.space_before = Pt(4)
p_cred.paragraph_format.space_after  = Pt(4)
add_run(p_cred,
        'OAB/PR 118.596   ·   Alta Performance Jurídica   ·   Sigilo Absoluto   ·   Atuação Nacional',
        font='Garamond', size=7.5, color=GOLD_RGB)

# ─────────────────────────────────────────────────────────────────────────────
# RODAPÉ (Word nativo)
# ─────────────────────────────────────────────────────────────────────────────

footer = sec.footer
for p in footer.paragraphs:
    p.clear()

ft = footer.add_table(rows=1, cols=1, width=Cm(21))
no_borders(ft)
set_table_width_pct(ft, 100)

fc = ft.rows[0].cells[0]
set_cell_bg(fc, DARK)
set_cell_margins(fc, top=200, left=500, bottom=200, right=500)
set_row_height(ft.rows[0], 900, exact=False)

for p in footer.paragraphs:
    p._element.getparent().remove(p._element)

# Nome no rodapé
p_fn = fc.paragraphs[0]
p_fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p_fn)

# Linha dourada no topo do rodapé
para_border_bottom(p_fn, color=GOLD, sz='6')
add_run(p_fn, '', font='Garamond', size=1)

p_f1 = fc.add_paragraph()
p_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p_f1)
p_f1.paragraph_format.space_before = Pt(4)
add_run(p_f1, 'JCPV  ADVOCACIA', font='Garamond', size=10,
        bold=True, color=GOLD_RGB)

p_f2 = fc.add_paragraph()
p_f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p_f2)
add_run(p_f2, 'www.jcpvadvocacia.com.br',
        font='Garamond', size=8, color=RGBColor(0xCC, 0xBB, 0x99))

# Rodapé com 3 colunas: esq / centro / dir
ft2 = footer.add_table(rows=1, cols=3, width=Cm(21))
no_borders(ft2)
set_table_width_pct(ft2, 100)

cols_data = [
    ('OAB/PR 118.596', WD_ALIGN_PARAGRAPH.LEFT),
    ('Sigilo Absoluto  ·  Segurança Jurídica', WD_ALIGN_PARAGRAPH.CENTER),
    ('Pág. ', WD_ALIGN_PARAGRAPH.RIGHT),
]

for i, (txt, align) in enumerate(cols_data):
    c = ft2.rows[0].cells[i]
    set_cell_bg(c, DARK)
    set_cell_margins(c, top=80, left=400, bottom=120, right=400)
    p = c.paragraphs[0]
    p.alignment = align
    no_space(p)
    add_run(p, txt, font='Garamond', size=7.5, color=GOLD_RGB)
    if i == 2:
        # Número de página automático
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_pg = p.add_run()
        run_pg.font.size = Pt(7.5)
        run_pg.font.color.rgb = GOLD_RGB
        run_pg.font.name = 'Garamond'
        run_pg._r.append(fldChar1)
        run_pg._r.append(instrText)
        run_pg._r.append(fldChar2)

# ─────────────────────────────────────────────────────────────────────────────
# CORPO DO DOCUMENTO — padrão petição ABNT
# ─────────────────────────────────────────────────────────────────────────────

def add_body_para(text='', align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  indent=True, bold=False, italic=False,
                  size=12, color=None, space_before=0, space_after=6,
                  underline=False):
    from docx.enum.text import WD_LINE_SPACING
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing      = Pt(18)   # 12pt × 1,5
    if text:
        r = add_run(p, text, size=size, bold=bold, italic=italic,
                    color=color or DARK_RGB, underline=underline)
    return p

# ── Destinatário ──────────────────────────────────────────────────────────────
p = add_body_para(
    'EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA '
    '___ª VARA CÍVEL DA COMARCA DE ________________________________',
    align=WD_ALIGN_PARAGRAPH.CENTER,
    indent=False, bold=True, size=12, space_after=24
)

# ── Qualificação ──────────────────────────────────────────────────────────────
p = add_body_para(
    '____________________________________ (nome completo do requerente), '
    'nacionalidade, estado civil, profissão, portador(a) do RG nº _____________ '
    'e CPF nº __________________, residente e domiciliado(a) na '
    '____________________________________________, nº ______, '
    'Bairro __________________, Cidade/UF, CEP ______________, '
    'por seu(sua) advogado(a) infra-assinado(a), vem respeitosamente à presença '
    'de Vossa Excelência propor a presente:',
    indent=True, space_after=12
)

# ── Título da ação ────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(6)
p_title.paragraph_format.space_after  = Pt(6)
p_title.paragraph_format.line_spacing = Pt(18)
add_run(p_title, 'AÇÃO DE __________________________________________',
        bold=True, size=12, color=DARK_RGB, underline=True)

# ── Em face de ────────────────────────────────────────────────────────────────
p_contra = doc.add_paragraph()
p_contra.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contra.paragraph_format.space_before = Pt(0)
p_contra.paragraph_format.space_after  = Pt(18)
p_contra.paragraph_format.line_spacing = Pt(18)
add_run(p_contra, 'em face de ___________________________________________',
        size=12, color=DARK_RGB)

# ── Seção I ───────────────────────────────────────────────────────────────────
def section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    add_run(p, text, bold=True, size=12, color=DARK_RGB)
    para_border_bottom(p, color=GOLD, sz='4')
    return p

section_title('I – DOS FATOS')

add_body_para(
    'Narra o(a) requerente que __________________________________________'
    '___________________________________________________________________'
    '____________________________________________________________, '
    'situação que gerou os prejuízos ora descritos e que fundamentam o '
    'presente pedido.',
    space_after=6
)
add_body_para(
    'Acrescenta ainda que _______________________________________________'
    '___________________________________________________________________'
    '______________________________________________, razão pela qual '
    'busca a tutela jurisdicional para proteção de seus direitos.',
    space_after=12
)

section_title('II – DO DIREITO')

add_body_para(
    'A pretensão deduzida encontra amparo na Lei nº _____________________, '
    'bem como nos princípios constitucionais da dignidade da pessoa humana '
    '(art. 1º, III, CF/88), da proteção ao mínimo existencial e nos '
    'demais dispositivos legais aplicáveis à espécie.',
    space_after=6
)
add_body_para(
    'Ademais, a jurisprudência pátria tem se posicionado no sentido de '
    '__________________________________________________________________'
    '________________________________, conforme reiteradas decisões do '
    'Colendo Superior Tribunal de Justiça.',
    space_after=12
)

section_title('III – DOS PEDIDOS')

add_body_para(
    'Diante do exposto, requer a Vossa Excelência que se digne a:',
    indent=True, space_after=6
)

pedidos = [
    'a) Receber e processar a presente ação, determinando a citação da parte requerida;',
    'b) Conceder os benefícios da justiça gratuita, nos termos do art. 98 do CPC/2015;',
    'c) Ao final, julgar TOTALMENTE PROCEDENTE o pedido, para que ________________________;',
    'd) Condenar a parte requerida ao pagamento das custas processuais e honorários advocatícios, '
       'nos termos do art. 85 do CPC/2015.',
]
for pedido in pedidos:
    add_body_para(pedido, indent=False,
                  space_before=2, space_after=4)

# ── Valor da causa ────────────────────────────────────────────────────────────
add_body_para('')
add_body_para(
    'Dá-se à causa o valor de R$ _________________ '
    '(___________________________________________).',
    indent=True, space_before=6, space_after=12
)

# ── Termos ────────────────────────────────────────────────────────────────────
add_body_para(
    'Nestes termos, pede e espera deferimento.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    indent=True, space_before=6, space_after=6
)

# Local e data
p_loc = doc.add_paragraph()
p_loc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_loc.paragraph_format.space_before = Pt(6)
p_loc.paragraph_format.space_after  = Pt(36)
p_loc.paragraph_format.line_spacing = Pt(18)
add_run(p_loc, 'Curitiba/PR, _____ de ________________________ de 20___.',
        size=12, color=DARK_RGB)

# ── Assinatura ────────────────────────────────────────────────────────────────
p_sig_line = doc.add_paragraph()
p_sig_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sig_line.paragraph_format.space_after = Pt(2)
add_run(p_sig_line, '________________________________________',
        size=12, color=GOLD_RGB)

p_sig_name = doc.add_paragraph()
p_sig_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sig_name.paragraph_format.space_after = Pt(2)
add_run(p_sig_name, 'Advogado(a) Responsável',
        bold=True, size=12, color=DARK_RGB)

p_sig_oab = doc.add_paragraph()
p_sig_oab.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sig_oab.paragraph_format.space_after = Pt(2)
add_run(p_sig_oab, 'OAB/PR 118.596',
        size=11, color=GOLD_RGB)

p_sig_firm = doc.add_paragraph()
p_sig_firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_sig_firm, 'JCPV ADVOCACIA  ·  www.jcpvadvocacia.com.br',
        size=10, color=GRAY_RGB)

# ─────────────────────────────────────────────────────────────────────────────
doc.save('/home/user/Apppainel-prestes/JCPV_Advocacia_Folha_Layout.docx')
print('Documento ABNT criado com sucesso!')
